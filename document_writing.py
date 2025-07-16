from docx import Document
import os
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

class DocumentLogger:
    def __init__(self, session_id, output_dir = "./chat_logs"):
        self.session_id = session_id
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.doc_path = os.path.join(output_dir, f"chat_session_{session_id}.docx")

        if os.path.exists(self.doc_path):
            self.document = Document(self.doc_path)
        else:
            self.document = Document()
            self.document.add_heading(f"Chat_Session_{session_id}", 0)
            self.document.save(self.doc_path)

    def write_interaction(self, question, response, sources = None):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        self.document.add_heading(f'User Question ({timestamp}):', level=2)

        self.document.add_paragraph(question)
        
        self.document.add_heading('AI Response:', level=2)
        self.document.add_paragraph(response)
        
        if sources:
            self.document.add_heading('Sources:', level=2)
            self.document.add_paragraph(sources)
        
        self.document.add_page_break()
        self.document.save(self.doc_path)

class DocumentWriter:
    def __init__(self, filename, output_dir = "./chat_logs"):
        self.filename = filename
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.doc_path = os.path.join(output_dir, f"{filename}.docx")
        self.structure_path = os.path.join(output_dir, f"{filename}_structure.json")


        if os.path.exists(self.doc_path):
            self.document = Document(self.doc_path)
        else:
            self.document = Document()
            title = self.document.add_heading(f"{filename}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.document.save(self.doc_path)

        if os.path.exists(self.structure_path):
            with open(self.structure_path, "r") as f:
                self.doc_structure = json.load(f)
        else:
            self.doc_structure = {}

        self.current_heading = None
        self.current_subheading = None

    def save_structure(self):
        with open(self.structure_path, "w") as f:
            json.dump(self.doc_structure, f, indent=2)

    def add_heading(self, heading_text):
        self.document.add_heading(heading_text, level=1)
        self.current_heading = heading_text
        self.current_subheading = None
        self.doc_structure.setdefault(heading_text, {})
        self.document.save(self.doc_path)
        self.save_structure()

    def add_subheading(self, subheading_text):
        if self.current_heading is None:
            print("Please add a heading first.")
            return
        self.document.add_heading(subheading_text, level=2)
        self.current_subheading = subheading_text
        self.doc_structure[self.current_heading].setdefault(subheading_text, [])
        self.document.save(self.doc_path)
        self.save_structure()

    def write_to_document(self, paragraph_text):
        self.document.add_paragraph(paragraph_text)
        if self.current_heading is None:
            self.doc_structure.setdefault("No Heading", []).append(paragraph_text)
        elif self.current_subheading:
            self.doc_structure[self.current_heading][self.current_subheading].append(paragraph_text)
        else:
            self.doc_structure[self.current_heading].setdefault("_content", []).append(paragraph_text)
        self.document.save(self.doc_path)
        self.save_structure()